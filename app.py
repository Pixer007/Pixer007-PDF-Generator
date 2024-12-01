import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime

# Function to load the Word template
def load_template(file_path):
    return Document(file_path)

# Function to replace placeholders in the template
def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value or key)  # Keep placeholder if value is missing

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value or key)

# Streamlit app
def main():
    st.title("Simple NDA Generator (Word File Only)")
    st.markdown("Generate a Word document by filling out the fields below.")

    # Sidebar input fields
    st.sidebar.header("Input Details")
    client_company_name = st.sidebar.text_input("Client Company Name", placeholder="Enter the company name")
    client_name_surname = st.sidebar.text_input("Client Full Name", placeholder="Enter the full name of the client")
    client_address = st.sidebar.text_area("Client Address", placeholder="Enter the client address")
    today_date = datetime.now().strftime("%d-%m-%Y")
    agreement_date = st.sidebar.text_input("Agreement Date", value=today_date)

    if st.sidebar.button("Generate NDA"):
        # Load the Word template
        template_path = "NDA Template - INDIA 3.docx"  # Ensure this file exists in the same directory
        doc = load_template(template_path)

        # Define replacements
        replacements = {
            "__Client Company Name (Client Name) __": client_company_name,
            "________": client_company_name,
            "__Client Address__": client_address,
            "10th September, 2023": agreement_date,
            "Name Surname": client_name_surname,
            "10-09-2023" : agreement_date,
        }

        # Replace placeholders in the Word document
        replace_placeholders(doc, replacements)

        # Save Word document to a BytesIO object
        word_output = BytesIO()
        doc.save(word_output)
        word_output.seek(0)

        # Allow the user to download the Word file
        st.success("NDA document generated successfully!")
        st.download_button(
            label="Download NDA (Word)",
            data=word_output,
            file_name=f"NDA_{client_company_name.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

if __name__ == "__main__":
    main()
